from tkinter import *
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
from time import sleep
from docx import Document
from docx.shared import Inches
from io import BytesIO
from one_piece_composite import one_piece_composite
from composite_in_section import composite_in_section
import cv2

class App:
    def __init__(self, root):
        self.root = root
        self.root.iconbitmap('App_Icons/icon2.ico')
        self.root.title('Анализ древесно-торфяного композита')
        self.root.geometry("900x600+300+150")
        self.root.config(bg='#E6E6E6')
        self.root.resizable(False, False)
        self.frames = {}
        self.current_image = None
        self.current_text = None

        self.images = {
            'main_bg': PhotoImage(file='App_Icons/Okno.png'),
            'form_bg': PhotoImage(file='App_Icons/form.png'),
            'okno2': PhotoImage(file='App_Icons/Okno2.png'),
            'okno3': PhotoImage(file='App_Icons/Okno3.png'),
            'okno4': PhotoImage(file='App_Icons/Okno4.png'),
            'okno5': PhotoImage(file='App_Icons/Okno5.png'),
            'b1': PhotoImage(file='App_Icons/b1.png'),
            'bind1': PhotoImage(file='App_Icons/bind1.png'),
            'b2': PhotoImage(file='App_Icons/b2.png'),
            'bind2': PhotoImage(file='App_Icons/bind2.png'),
            'b3': PhotoImage(file='App_Icons/b3.png'),
            'bind3': PhotoImage(file='App_Icons/bind3.png'),
            'b4': PhotoImage(file='App_Icons/b4.png'),
            'bind4': PhotoImage(file='App_Icons/bind4.png'),
            'bf1': PhotoImage(file='App_Icons/bf1.png'),
            'bindfb1': PhotoImage(file='App_Icons/bindfb1.png'),
            'bf2': PhotoImage(file='App_Icons/bf2.png'),
            'bindfb2': PhotoImage(file='App_Icons/bindfb2.png'),

        }

        self.show_frame("main")

    def show_frame(self, name):
        for f in self.frames.values():
            f.pack_forget()

        if name not in self.frames:
            frame_creator = getattr(self, f"create_{name}_frame")
            self.frames[name] = frame_creator()
        self.frames[name].pack(fill='both', expand=True)

    def load_and_analyze1(self):
        image_path = filedialog.askopenfilename(filetypes=[("Image files", "*.png")], title='Выберите изображение')
        if not image_path:
            return

        preview_canvas = None

        try:
            # Canvas для превью (смещен правее и ниже)
            preview_canvas = Canvas(self.frames['analysis1'], width=364, height=330, bg='#eeeeee',
                                         highlightthickness=0, bd=2, relief='solid')
            preview_canvas.place(x=476, y=159)  # Смещение координат

            # 1. Показываем превью
            pil_image = Image.open(image_path)
            pil_image.thumbnail((700, 700), Image.LANCZOS)
            self.current_image = ImageTk.PhotoImage(pil_image)

            if preview_canvas:
                preview_canvas.delete("all")
                preview_canvas.create_image(180, 125, image=self.current_image)

            # 2. Выполняем анализ (имитация)
            root.update()  # Обновляем интерфейс перед долгой операцией
            sleep(2)  # Имитация анализа

            # 3. Переходим к результатам
            processed_image, result_text = one_piece_composite(image_path)

            # Конвертация из BGR в RGB
            processed_image_rgb = cv2.cvtColor(processed_image.astype('uint8'), cv2.COLOR_BGR2RGB)

            # Создание изображения для Tkinter
            result_image = Image.fromarray(processed_image_rgb)
            result_image.thumbnail((700, 700), Image.LANCZOS)
            result_image = ImageTk.PhotoImage(result_image)

            self.current_image = result_image
            self.current_text = result_text
            preview_canvas.destroy()

            self.show_frame('result1')

            # 4. Показываем результат
            result_canvas = Canvas(self.frames['result1'], width=364, height=330, bg='#eeeeee',
                                   highlightthickness=0, bd=2, relief='solid')
            result_canvas.place(x=476, y=159)  # Смещение координат
            if result_canvas:
                result_canvas.delete("all")
                result_canvas.create_image(180, 125, image=self.current_image)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить изображение:\n{str(e)}")
            if preview_canvas:
                preview_canvas.destroy()

    def load_and_analyze2(self):
        image_path = filedialog.askopenfilename(filetypes=[("Image files", "*.png")],
                                                title='Выберите изображение')
        if not image_path:
            return

        preview_canvas = None

        try:
            # Canvas для превью (смещен правее и ниже)
            preview_canvas = Canvas(self.frames['analysis2'], width=364, height=330, bg='#eeeeee',
                                         highlightthickness=0, bd=2, relief='solid')
            preview_canvas.place(x=476, y=159)  # Смещение координат

            # 1. Показываем превью
            pil_image = Image.open(image_path)
            pil_image.thumbnail((700, 700), Image.LANCZOS)
            self.current_image = ImageTk.PhotoImage(pil_image)

            if preview_canvas:
                preview_canvas.delete("all")
                preview_canvas.create_image(180, 125, image=self.current_image)

            # 2. Выполняем анализ (имитация)
            root.update()  # Обновляем интерфейс перед долгой операцией
            sleep(1)  # Имитация анализа

            # 3. Переходим к результатам
            processed_image, result_text = composite_in_section(image_path)
            result_image = Image.fromarray(processed_image.astype('uint8'), 'RGB')
            result_image.thumbnail((700, 700), Image.LANCZOS)
            result_image = ImageTk.PhotoImage(result_image)

            self.current_image = result_image
            preview_canvas.destroy()

            self.show_frame('result2')

            # 4. Показываем результат
            result_canvas = Canvas(self.frames['result2'], width=364, height=330, bg='#eeeeee',
                                   highlightthickness=0, bd=2, relief='solid')
            result_canvas.place(x=476, y=159)  # Смещение координат
            if result_canvas:
                result_canvas.delete("all")
                result_canvas.create_image(180, 125, image=self.current_image)

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить изображение:\n{str(e)}")
            if preview_canvas:
                preview_canvas.destroy()

    @staticmethod
    def save_file(image, text):
        try:
            pil_image = ImageTk.getimage(image)
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx")],
                title="Сохранить отчет"
            )
            if file_path:
                # Создаем документ Word
                doc = Document()

                # Добавляем текст
                doc.add_paragraph("Отчет анализа композита")

                # Сохраняем изображение во временный файл
                img_stream = BytesIO()
                pil_image.save(img_stream, format='PNG')
                img_stream.seek(0)

                # Добавляем изображение в Word
                doc.add_picture(img_stream, width=Inches(4.0))  # Ширина 4 дюйма

                doc.add_paragraph(text)
                doc.save(file_path)
                messagebox.showinfo("Успех", "Отчет успешно сохранен!")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить отчет:\n{str(e)}")

    def create_main_frame(self):
        frame = Frame(self.root, bg='#E6E6E6')
        Label(frame, image=self.images['main_bg']).place(x=-1, y=0)

        b1 = Button(frame, image=self.images['b1'], bg='#E6E6E6', borderwidth=0,
                    highlightthickness=0, activebackground='#E6E6E6',
                    command=lambda: self.show_frame("form"))
        b1.place(x=460, y=134)
        b1.bind('<Enter>', lambda e: b1.config(image=self.images['bind1']))
        b1.bind('<Leave>', lambda e: b1.config(image=self.images['b1']))
        return frame

    def create_form_frame(self):
        frame = Frame(self.root, bg='#E6E6E6')
        Label(frame, image=self.images['form_bg']).place(x=-1, y=0)

        bf1 = Button(frame, image=self.images['bf1'], bg='#EDEDED', borderwidth=0,
                         highlightthickness=0, activebackground='#EDEDED',
                         command=lambda: self.show_frame("analysis1"))
        bf1.place(x=480, y=186)
        bf1.bind('<Enter>', lambda e: bf1.config(image=self.images['bindfb1']))
        bf1.bind('<Leave>', lambda e: bf1.config(image=self.images['bf1']))

        bf2 = Button(frame, image=self.images['bf2'], bg='#EDEDED', borderwidth=0,
                         highlightthickness=0, activebackground='#EDEDED',
                         command=lambda: self.show_frame("analysis2"))
        bf2.place(x=480, y=293)
        bf2.bind('<Enter>', lambda e: bf2.config(image=self.images['bindfb2']))
        bf2.bind('<Leave>', lambda e: bf2.config(image=self.images['bf2']))

        # 👇 Добавленная кнопка аналог b22
        b22 = Button(frame, image=self.images['b2'], bg='#E6E6E6', borderwidth=0,
                         highlightthickness=0, activebackground='#E6E6E6',
                         command=lambda: self.show_frame("main"))
        b22.place(x=797, y=47)
        b22.bind('<Enter>', lambda e: b22.config(image=self.images['bind2']))
        b22.bind('<Leave>', lambda e: b22.config(image=self.images['b2']))

        return frame

    def create_analysis1_frame(self):
        return self.create_analysis_frame('okno2', lambda: self.load_and_analyze1())

    def create_analysis2_frame(self):
        return self.create_analysis_frame('okno5', lambda: self.load_and_analyze2())

    def create_analysis_frame(self, bg_name, next_command):
        frame = Frame(self.root, bg='#E6E6E6')
        Label(frame, image=self.images[bg_name]).place(x=-1, y=0)

        b2 = Button(frame, image=self.images['b2'], bg='#E6E6E6', borderwidth=0,
                        highlightthickness=0, activebackground='#E6E6E6',
                        command=lambda: self.show_frame("form"))
        b2.place(x=797, y=47)
        b2.bind('<Enter>', lambda e: b2.config(image=self.images['bind2']))
        b2.bind('<Leave>', lambda e: b2.config(image=self.images['b2']))

        b3 = Button(frame, image=self.images['b3'], bg='#E6E6E6', borderwidth=0,
                        highlightthickness=0, activebackground='#E6E6E6',
                        command=next_command)
        b3.place(x=645, y=515)
        b3.bind('<Enter>', lambda e: b3.config(image=self.images['bind3']))
        b3.bind('<Leave>', lambda e: b3.config(image=self.images['b3']))
        return frame

    def create_result1_frame(self):
        return self.create_result_frame('okno3')

    def create_result2_frame(self):
        return self.create_result_frame('okno4')

    def create_result_frame(self, bg_name):
        frame = Frame(self.root, bg='#E6E6E6')
        Label(frame, image=self.images[bg_name]).place(x=-1, y=0)

        b2 = Button(frame, image=self.images['b2'], bg='#E6E6E6', borderwidth=0,
                        highlightthickness=0, activebackground='#E6E6E6',
                        command=lambda: self.show_frame("form"))
        b2.place(x=797, y=47)
        b2.bind('<Enter>', lambda e: b2.config(image=self.images['bind2']))
        b2.bind('<Leave>', lambda e: b2.config(image=self.images['b2']))

        b4 = Button(frame, image=self.images['b4'], bg='#E6E6E6', borderwidth=0,
                        highlightthickness=0, activebackground='#E6E6E6',
                    command=lambda: self.save_file(self.current_image, self.current_text))
        b4.place(x=645, y=515)
        b4.bind('<Enter>', lambda e: b4.config(image=self.images['bind4']))
        b4.bind('<Leave>', lambda e: b4.config(image=self.images['b4']))
        return frame

if __name__ == "__main__":
    root = Tk()
    app = App(root)
    root.mainloop()
